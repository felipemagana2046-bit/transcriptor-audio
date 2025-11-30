import streamlit as st
import google.generativeai as genai
from docx import Document
import os
import tempfile
from datetime import datetime

# --- FUNCIÓN MÁGICA PARA PONER NEGRITAS EN WORD ---
def add_markdown_paragraph(doc, text):
    """
    Convierte texto con **negritas** estilo markdown a negritas reales de Word.
    """
    if not text:
        return

    # Si hay saltos de línea dobles, crear párrafos nuevos
    paragraphs = text.split('\n\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
            
        p = doc.add_paragraph()
        # Separar por los asteriscos dobles
        parts = para_text.split('**')
        
        # El truco: Las partes impares (1, 3, 5...) son las que estaban entre asteriscos
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 != 0: # Si es impar, es negrita real de Word
                run.bold = True

# 1. Configuración de la Página
st.set_page_config(
    page_title="Transcriptor Pro",
    page_icon="🎙️",
    layout="centered"
)

# 2. Seguridad y Configuración de API Key
api_key = None
# Intentamos leer de secrets.toml
if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    # Fallback manual
    with st.sidebar:
        st.header("Configuración")
        api_key = st.text_input("Ingresa tu Gemini API Key", type="password")

if api_key:
    genai.configure(api_key=api_key)

# 3. Interfaz de Usuario (UI)
st.title("🎙️ Transcriptor de Audio")

# Configuración (SIN FECHA)
with st.expander("Configuración de Interlocutores", expanded=True):
    main_speaker = st.text_input("Nombre del Interlocutor Principal (La voz dominante)", value="La Madre")
    other_speakers = st.text_input("Otros interlocutores (opcional)", placeholder="Ej: Pedro, La Abuela, El Doctor")

# Archivo
st.subheader("Archivo de Audio")
uploaded_file = st.file_uploader("Sube tu audio", type=["mp3", "wav", "m4a", "ogg", "aac", "flac"])

# Botón de Acción
if st.button("Transcribir y Generar Word", type="primary"):
    if not api_key:
        st.error("⚠️ Falta la API Key.")
    elif not uploaded_file:
        st.error("⚠️ Por favor sube un archivo de audio.")
    else:
        try:
            with st.spinner("Procesando audio... Esto puede tardar unos minutos."):
                # Guardar archivo temporalmente
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_file_path = tmp_file.name

                # Subir a Gemini
                st.text("Subiendo audio a Gemini...")
                myfile = genai.upload_file(tmp_file_path)
                
                # Construcción del Prompt Mejorado
                speaker_instructions = f"Identifica a la voz principal/dominante y etiquetarla como **{main_speaker}**."
                if other_speakers:
                    speaker_instructions += f" Si detectas otras voces, usa los nombres: {other_speakers}."
                
                system_prompt = f"""
                Actúa como un editor experto. Tu tarea es transcribir este audio LITERALMENTE, pero cuidando mucho la legibilidad.
                
                REGLAS DE FORMATO:
                1. {speaker_instructions}
                2. IMPORTANTE: Usa **negritas** (asteriscos dobles) para los nombres (ej: **{main_speaker}:** ...).
                3. DIVIDE EL TEXTO EN PÁRRAFOS CORTOS. No hagas bloques gigantes de texto.
                4. Usa un doble salto de línea cada vez que cambie la idea o el interlocutor.
                """

                # Selección Inteligente de Modelo
                file_size_mb = uploaded_file.size / (1024 * 1024)
                
                # Lógica de Enrutamiento
                if file_size_mb > 50:
                    model_name = "gemini-2.5-flash"
                    st.info(f"📂 Archivo grande ({file_size_mb:.1f} MB). Usando Gemini 2.5 Flash para velocidad.")
                else:
                    model_name = "gemini-3-pro-preview"
                    st.success(f"🧠 Archivo ligero ({file_size_mb:.1f} MB). Usando Gemini 3 Pro para inteligencia.")

                # Generar contenido con manejo de errores
                st.text(f"Transcribiendo con {model_name}...")
                
                try:
                    model = genai.GenerativeModel(model_name)
                    response = model.generate_content([myfile, system_prompt])
                    transcription_text = response.text
                except Exception as e:
                    # Fallback automático
                    if model_name == "gemini-3-pro-preview":
                        st.warning("⚠️ El modelo Pro está ocupado, cambiando a Flash automáticamente...")
                        model = genai.GenerativeModel("gemini-2.5-flash")
                        response = model.generate_content([myfile, system_prompt])
                        transcription_text = response.text
                    else:
                        raise e

                # --- CREAR DOCX CON NEGRITAS REALES ---
                doc = Document()
                doc.add_heading(uploaded_file.name, 0) # Título del archivo
                doc.add_paragraph("--- Transcripción Generada por IA ---")
                
                # Usamos la función especial
                add_markdown_paragraph(doc, transcription_text)
                
                # Guardar DOCX temporalmente
                docx_filename = f"Transcripcion_{uploaded_file.name.rsplit('.', 1)[0]}.docx"
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                    doc.save(tmp_docx.name)
                    tmp_docx_path = tmp_docx.name
                
                # Leer para descarga
                with open(tmp_docx_path, "rb") as f:
                    docx_data = f.read()

                st.success("¡Transcripción completada!")
                
                # Botón de Descarga
                st.download_button(
                    label="📥 Descargar Transcripción (.docx)",
                    data=docx_data,
                    file_name=docx_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # Limpieza
                os.unlink(tmp_file_path)
                os.unlink(tmp_docx_path)

        except Exception as e:

            st.error(f"Ocurrió un error: {str(e)}")
